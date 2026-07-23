"""
core/structural_check.py — mechanism-agnostic manuscript verification.

Reads a .docx directly via python-docx (already a dependency) and checks
the things that actually matter regardless of what produced the file:
does every expected section exist, in the right order; does the blinded
version actually contain no identifying strings.

Deliberately NOT built on build.js. See project discussion: build.js's
only real justification is as a safe surface for iterative str_replace
editing — it has no bearing on verification, and reconstructing XML into
a synthetic build.js purely to re-check it would reintroduce the exact
run-splitting fragility this design avoids, for a one-shot check that
never gets edited. This module is the one verification path used
everywhere: this app's own Step 5/6 gate, Admin's manual re-check, and
"hired gun" mode on a third-party docx that has never seen this tool.

What this deliberately does NOT do (state these limits to the user, not
just here): does not validate OOXML well-formedness or file corruption at
the format level (a genuinely malformed docx needs Word's own "Open and
Repair" or a dedicated OOXML validator, not this). Does not check docx
*metadata* (author/company core properties, revision history) for
identifying content — Word's own "Inspect Document" covers that; this
only checks visible body text.
"""
from docx import Document


HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title"}


def _heading_paragraphs(doc):
    """Returns [(style_name, text), ...] for every heading-styled paragraph, in document order."""
    return [
        (p.style.name, p.text.strip())
        for p in doc.paragraphs
        if p.style and p.style.name in HEADING_STYLES and p.text.strip()
    ]


def check_structure(docx_path, expected_sections):
    """
    expected_sections: list of section names in intended order (e.g. from
    the regime's sections list — ["Abstract", "Introduction", ...]).

    Returns (ok: bool, findings: list[str]). findings is empty iff ok.
    Mirrors structural_lint.js's intent (section presence + order), not
    its exact mechanism — there is no build.js-specific equivalent here
    for the orphaned-function-body check, since that's a source-editing
    bug signature with no rendered-document equivalent; a section that
    silently never got written just shows up as "missing" below, which
    this already catches, just without the more precise "here's why".
    """
    findings = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        return False, [f"Could not open {docx_path} as a .docx: {e}"]

    headings = _heading_paragraphs(doc)
    heading_texts = [text for _, text in headings]

    # Missing: every expected section name should appear as some heading's
    # text (exact match, since this app always writes section names
    # verbatim as Heading 1 text — see core/docx_export.py).
    missing = [s for s in expected_sections if s not in heading_texts]
    for m in missing:
        findings.append(f"Expected section \"{m}\" not found as a heading anywhere in the document.")

    # Order: among the expected sections that ARE present, their relative
    # order in the document should match expected_sections' order.
    present_expected = [s for s in expected_sections if s in heading_texts]
    actual_order = [t for t in heading_texts if t in present_expected]
    if actual_order != present_expected:
        findings.append(
            "Section order does not match the regime's intended order. "
            f"Expected: {present_expected} — Found: {actual_order}"
        )

    # Unexpected: headings that don't correspond to any expected section
    # name — not necessarily wrong (front matter, subsections), so this is
    # informational, not a failure on its own.
    unexpected = [t for t in heading_texts if t not in expected_sections]
    if unexpected:
        findings.append(
            f"Note: {len(unexpected)} heading(s) found that aren't in the expected section "
            f"list (may be intentional subsections/front matter): {unexpected}"
        )
        # This one is informational only — don't let it fail the check by itself.
        hard_findings = [f for f in findings if not f.startswith("Note:")]
        return (len(hard_findings) == 0), findings

    return (len(findings) == 0), findings


def check_blind(docx_path, identifying_strings):
    """
    Returns (ok: bool, findings: list[str]). Reads every paragraph's
    actual text (python-docx already resolves split runs into one string
    per paragraph — the same problem raw XML parsing would have to solve
    by hand) and checks none of the declared identifying strings appear
    anywhere in the visible body text.

    Does NOT check docx core properties (author/company metadata) — see
    module docstring. Does NOT check headers/footers/footnotes/comments —
    a real gap worth knowing about, not silently assumed to be covered.
    """
    findings = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        return False, [f"Could not open {docx_path} as a .docx: {e}"]

    full_text = "\n".join(p.text for p in doc.paragraphs)

    for s in identifying_strings:
        s = s.strip()
        if not s:
            continue
        if s in full_text:
            count = full_text.count(s)
            findings.append(f'Identifying string "{s}" found {count} time(s) in the document body.')

    return (len(findings) == 0), findings
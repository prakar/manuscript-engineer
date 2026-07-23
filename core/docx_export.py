"""
core/docx_export.py — renders a manuscript's sections into a .docx
scaffold. No Streamlit dependency; testable standalone.
"""
from docx import Document as DocxDocument


def export_manuscript(manuscript, sections, footer_text, blind=False,
                       identifying_strings=None):
    """
    manuscript: dict from db.get_manuscript()
    sections: list of dicts from db.get_sections()
    Returns a docx.Document. Caller saves it wherever appropriate.

    If blind=True, a best-effort strip of any identifying string is run
    against every section's content before writing — this is a scaffold
    convenience, NOT a substitute for blind_check_template.js-style
    verification against the real rendered artifact. That distinction is
    stated in the generated document's colophon, not silently assumed.
    """
    doc = DocxDocument()

    identifying_strings = identifying_strings or []

    def _redact(text):
        if not blind:
            return text
        for s in identifying_strings:
            if s and s in text:
                text = text.replace(s, "[REDACTED]")
        return text

    doc.add_heading(_redact(manuscript.get("title") or "[[TODO: title]]"), level=0)
    doc.add_paragraph(f"Regime statement: {_redact(manuscript.get('regime_statement', ''))}")
    doc.add_paragraph(f"Central claim: {_redact(manuscript.get('claim', ''))}")

    for sec in sections:
        doc.add_heading(sec["name"], level=1)
        content = sec.get("content") or f"[[TODO: draft {sec['name']} content]]"
        if sec.get("hint"):
            doc.add_paragraph(f"[[Drafting note: {sec['hint']}]]").italic = True
        content = _redact(content)
        doc.add_paragraph(content)

    doc.add_page_break()
    colophon = (
        "This document is a structured DRAFT, not a finished or "
        "reviewed manuscript. "
    )
    if blind:
        colophon += (
            "Name/institution redaction above is a best-effort convenience "
            "only — before real submission, verify the actual document with "
            "blind_check_template.js, not this note. "
        )
    colophon += footer_text
    p = doc.add_paragraph(colophon)
    p.italic = True

    return doc